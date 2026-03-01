#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档解析器 - 支持 Word 和 PDF 文档解析
"""

import os
from pathlib import Path
from typing import Optional


def parse_word(file_path: str) -> str:
    """
    解析 Word 文档，提取纯文本
    
    Args:
        file_path: Word 文档路径
        
    Returns:
        提取的文本内容
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    doc = Document(file_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # 处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    
    return "\n".join(paragraphs)


def parse_pdf(file_path: str) -> str:
    """
    解析 PDF 文档，提取纯文本
    
    Args:
        file_path: PDF 文档路径
        
    Returns:
        提取的文本内容
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("请安装 pdfplumber: pip install pdfplumber")
    
    import unicodedata
    
    text_parts = []
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                # Unicode 规范化：将 CJK 兼容字符转换为标准字符
                text = unicodedata.normalize('NFKC', text)
                text_parts.append(text.strip())
    
    return "\n".join(text_parts)


def parse_document(file_path: str) -> str:
    """
    根据文件类型自动选择解析器
    
    Args:
        file_path: 文档路径（支持 .docx, .pdf, .txt）
        
    Returns:
        提取的文本内容
        
    Raises:
        ValueError: 不支持的文件格式
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    suffix = path.suffix.lower()
    
    if suffix == ".docx":
        return parse_word(file_path)
    elif suffix == ".pdf":
        return parse_pdf(file_path)
    elif suffix == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"不支持的文件格式: {suffix}（支持 .docx, .pdf, .txt）")


def get_supported_files(directory: str) -> list:
    """
    获取目录下所有支持的文件
    
    Args:
        directory: 目录路径
        
    Returns:
        支持的文件路径列表
    """
    supported_extensions = {".docx", ".pdf", ".txt"}
    files = []
    
    for path in Path(directory).iterdir():
        if path.is_file() and path.suffix.lower() in supported_extensions:
            files.append(str(path))
    
    return sorted(files)


if __name__ == "__main__":
    # 测试代码
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            text = parse_document(file_path)
            print(f"成功解析文件: {file_path}")
            print(f"文本长度: {len(text)} 字符")
            print("-" * 50)
            print(text[:500] + "..." if len(text) > 500 else text)
        except Exception as e:
            print(f"解析失败: {e}")
    else:
        print("用法: python parse_document.py <文件路径>")
