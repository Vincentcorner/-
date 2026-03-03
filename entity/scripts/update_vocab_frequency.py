# -*- coding: utf-8 -*-
"""
词频统计脚本 - 直接更新原标准词库的词频列
"""

import pandas as pd
from docx import Document
import re
import os

def read_docx(file_path):
    """读取docx文件，返回全部文本内容"""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 也读取表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

def count_word_in_text(text, word):
    """统计单个词在文本中出现的次数"""
    if pd.isna(word) or not word:
        return 0
    word = str(word).strip()
    if not word:
        return 0
    try:
        return len(re.findall(re.escape(word), text))
    except:
        return 0

def main():
    # 文件路径
    vocab_path = r'D:\数字研究院工作\认知世界大模型\第三个路径\entity\entity\result\标准词库.xlsx'
    docx_path = r'D:\数字研究院工作\认知世界大模型\第三个路径\entity\entity\originalfile\广东省失业保险条例.docx'
    
    print("=" * 50)
    print("词频统计工具 - 更新原词库")
    print("=" * 50)
    
    # 检查文件是否存在
    print(f"\n1. 检查文件...")
    if not os.path.exists(vocab_path):
        print(f"   错误：找不到标准词库文件: {vocab_path}")
        return
    
    if not os.path.exists(docx_path):
        print(f"   错误：找不到原始文档: {docx_path}")
        return
    
    print(f"   标准词库: 存在")
    print(f"   原始文档: 存在")
    
    # 读取标准词库
    print(f"\n2. 读取标准词库...")
    df_vocab = pd.read_excel(vocab_path)
    print(f"   列名: {df_vocab.columns.tolist()}")
    print(f"   词汇数量: {len(df_vocab)}")
    
    # 识别词汇列（通常是"实体名"或第一列）
    word_column = None
    for col in ['实体名', '词汇', '标准词']:
        if col in df_vocab.columns:
            word_column = col
            break
    
    if word_column is None:
        word_column = df_vocab.columns[0]
    
    print(f"   使用词汇列: {word_column}")
    
    # 读取原始文档
    print(f"\n3. 读取原始文档...")
    doc_text = read_docx(docx_path)
    print(f"   文档字符数: {len(doc_text)}")
    
    # 统计词频并更新
    print(f"\n4. 统计词频...")
    
    # 确保有词频列
    if '词频' not in df_vocab.columns:
        df_vocab['词频'] = 0
        print("   已创建词频列")
    
    # 逐行统计词频
    for idx in range(len(df_vocab)):
        word = df_vocab.loc[idx, word_column]
        freq = count_word_in_text(doc_text, word)
        df_vocab.loc[idx, '词频'] = freq
        
        if (idx + 1) % 50 == 0:
            print(f"   已处理 {idx + 1}/{len(df_vocab)} 个词汇...")
    
    # 统计结果
    total_words = len(df_vocab)
    words_with_freq = len(df_vocab[df_vocab['词频'] > 0])
    words_without_freq = len(df_vocab[df_vocab['词频'] == 0])
    
    print(f"\n5. 统计结果:")
    print(f"   总词汇数: {total_words}")
    print(f"   有词频的词汇数: {words_with_freq}")
    print(f"   无词频的词汇数: {words_without_freq}")
    
    # 保存回原文件
    print(f"\n6. 保存结果...")
    df_vocab.to_excel(vocab_path, index=False)
    print(f"   已更新标准词库: {vocab_path}")
    
    # 显示前10个高频词
    print(f"\n7. 前10个高频词汇:")
    df_sorted = df_vocab.sort_values('词频', ascending=False).head(10)
    for _, row in df_sorted.iterrows():
        print(f"   {row[word_column]}: {row['词频']} 次")
    
    print("\n" + "=" * 50)
    print("词频统计完成！已更新原标准词库文件。")
    print("=" * 50)

if __name__ == '__main__':
    main()
